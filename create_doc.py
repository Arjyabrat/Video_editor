"""Generate a professional DOCX documentation file for the Video Editor project."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_documentation():
    doc = Document()
    
    # Configure styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Video Editor', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(36)
        run.font.color.rgb = RGBColor(35, 134, 54)
    
    # Subtitle
    subtitle = doc.add_paragraph('Automated Screen Recording Trimmer')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    subtitle.runs[0].font.italic = True
    
    doc.add_paragraph()
    
    # Overview Section
    doc.add_heading('Overview', level=1)
    overview = doc.add_paragraph()
    overview.add_run('Video Editor').bold = True
    overview.add_run(' is a FastAPI-based web application that automatically trims long idle or static sections from screen recordings while preserving active content. It provides an intuitive browser-based interface for uploading, processing, previewing, and downloading optimized videos.')
    
    # Key Features
    doc.add_heading('Key Features', level=1)
    features = [
        ('Browser-Based Upload', 'Upload and process video files directly from your web browser'),
        ('Smart Detection', 'Automatic static-scene detection using SSIM, motion analysis, and pixel-change checks'),
        ('Preview Player', 'Built-in video player with 5-second skip controls for reviewing results'),
        ('Export Options', 'YouTube-style menu for selecting playback speed and quality settings'),
        ('Instant Download', 'Download processed output files with a single click'),
    ]
    
    for title_text, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title_text + ': ').bold = True
        p.add_run(desc)
    
    # Tech Stack
    doc.add_heading('Technology Stack', level=1)
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Component', 'Technology']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '238636')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    tech_data = [
        ('Backend Framework', 'FastAPI'),
        ('Video Analysis', 'OpenCV + NumPy'),
        ('Video Rendering', 'FFmpeg (via imageio-ffmpeg)'),
        ('Frontend', 'Embedded HTML/CSS/JavaScript'),
    ]
    
    for i, (component, tech) in enumerate(tech_data, 1):
        table.rows[i].cells[0].text = component
        table.rows[i].cells[1].text = tech
    
    doc.add_paragraph()
    
    # Project Structure
    doc.add_heading('Project Structure', level=1)
    
    structure = doc.add_paragraph()
    structure.paragraph_format.left_indent = Inches(0.25)
    code_style = structure.add_run('''Video-Editor/
├── src/
│   └── video_editor/
│       ├── __init__.py
│       ├── app.py            # FastAPI app + UI + API routes
│       ├── analysis.py       # Static/idle segment detection
│       ├── ffmpeg_tools.py   # FFmpeg render pipeline
│       └── service.py        # Job orchestration
├── storage/
│   ├── uploads/              # Uploaded source videos
│   └── outputs/              # Generated output videos
├── requirements.txt
└── README.md''')
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(10)
    
    # Installation
    doc.add_heading('Installation', level=1)
    
    doc.add_heading('Requirements', level=2)
    reqs = ['Python 3.10 or higher', 'FFmpeg (handled automatically by imageio-ffmpeg)', 'Windows, macOS, or Linux']
    for req in reqs:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading('Setup Steps', level=2)
    steps = [
        'Clone or download the repository',
        'Open a terminal in the project directory',
        'Install dependencies: pip install -r requirements.txt',
        'Run the application (see Usage section)',
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. ').bold = True
        p.add_run(step)
    
    # Usage
    doc.add_heading('Usage', level=1)
    
    doc.add_heading('Starting the Server', level=2)
    cmd = doc.add_paragraph()
    cmd.add_run('python -m uvicorn src.video_editor.app:app --host 127.0.0.1 --port 8000')
    cmd.runs[0].font.name = 'Consolas'
    cmd.runs[0].font.size = Pt(10)
    cmd.paragraph_format.left_indent = Inches(0.25)
    
    access = doc.add_paragraph()
    access.add_run('Then open your browser and navigate to: ')
    access.add_run('http://127.0.0.1:8000').bold = True
    
    doc.add_heading('Processing Workflow', level=2)
    workflow = [
        'Upload a screen recording (MP4, MOV, MKV, etc.)',
        'Configure detection options (static seconds, sample FPS, thresholds)',
        'Click "Process Video" to start analysis',
        'Backend analyzes static spans and compresses long static sections',
        'Preview the rendered output in the built-in player',
        'Download the processed video',
    ]
    for i, step in enumerate(workflow, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. ').bold = True
        p.add_run(step)
    
    # Processing Settings
    doc.add_heading('Processing Settings', level=1)
    
    settings_table = doc.add_table(rows=9, cols=3)
    settings_table.style = 'Table Grid'
    
    settings_headers = ['Setting', 'Default', 'Description']
    for i, header in enumerate(settings_headers):
        cell = settings_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '238636')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    settings_data = [
        ('static_seconds', '5.0', 'Minimum idle duration before trimming applies'),
        ('sample_fps', '3.0', 'Analysis sampling rate (frames per second)'),
        ('ssim_threshold', '0.98', 'Visual similarity threshold (0-1)'),
        ('motion_threshold', '0.45', 'Global motion tolerance'),
        ('cursor_threshold', '3.5', 'Pointer/feature motion tolerance'),
        ('min_keep_seconds', '0.25', 'Minimum output segment length'),
        ('export_speed', '1.0', 'Output playback speed multiplier'),
        ('export_quality', '720p', 'Output quality profile (480p, 720p, 1080p)'),
    ]
    
    for i, (setting, default, desc) in enumerate(settings_data, 1):
        settings_table.rows[i].cells[0].text = setting
        settings_table.rows[i].cells[0].paragraphs[0].runs[0].font.name = 'Consolas'
        settings_table.rows[i].cells[1].text = default
        settings_table.rows[i].cells[2].text = desc
    
    doc.add_paragraph()
    
    # API Reference
    doc.add_heading('API Reference', level=1)
    
    endpoints = [
        ('POST /api/process', 'Upload video and settings, run detection and rendering, return output metadata'),
        ('POST /api/reexport', 'Re-render existing upload with new speed/quality settings'),
        ('GET /api/download/{output_name}', 'Download the generated MP4 file'),
    ]
    
    for endpoint, desc in endpoints:
        p = doc.add_paragraph()
        p.add_run(endpoint).bold = True
        p.add_run(f'\n{desc}')
    
    # How It Works
    doc.add_heading('How Static Detection Works', level=1)
    
    detection_info = doc.add_paragraph()
    detection_info.add_run('The detection algorithm analyzes video frames to identify static (idle) sections:\n\n')
    
    bullets = [
        'Static spans are detected when frames remain highly similar for at least the configured threshold',
        'Long static parts are compressed, not fully removed, keeping 2-3 seconds of context',
        'Small UI updates (like typing) are treated conservatively to avoid over-cutting',
        'Multiple metrics are used: SSIM similarity, motion vectors, and pixel delta analysis',
    ]
    
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')
    
    # Troubleshooting
    doc.add_heading('Troubleshooting', level=1)
    
    issues = [
        ('Port Already in Use', 'Run: Get-NetTCPConnection -LocalPort 8000 -State Listen\nThen: Stop-Process -Id <PID> -Force'),
        ('Processing Takes Too Long', 'Reduce sample_fps (e.g., from 3 to 2), keep static_seconds around 5, use 720p export'),
        ('Upload Not Processing', 'Refresh the browser page and retry the upload'),
    ]
    
    for issue, solution in issues:
        p = doc.add_paragraph()
        p.add_run(issue + ': ').bold = True
        p.add_run(solution)
        doc.add_paragraph()
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('─' * 50)
    doc.add_paragraph()
    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer2.add_run('Video Editor Documentation').italic = True
    footer2.add_run(' | Generated with Python-docx')
    
    # Save
    doc.save('Video_Editor_Documentation.docx')
    print('Documentation created: Video_Editor_Documentation.docx')

if __name__ == '__main__':
    create_documentation()
